import io
import json
import re
from typing import List

import httpx
from docx import Document
from fastapi import FastAPI, File, Form, HTTPException, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response, StreamingResponse
from pydantic import BaseModel

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Contact info detection
CONTACT_RE = re.compile(
    r"[\w.+-]+@[\w-]+\.[a-zA-Z]{2,}"         # email
    r"|(?:\+?1[-.\s]?)?\(?\d{3}\)?[-.\s]\d{3}[-.\s]\d{4}"  # phone
    r"|linkedin\.com"                          # LinkedIn
    r"|github\.com"                            # GitHub
    r"|gitlab\.com"
    r"|portfolio|website|http",
    re.IGNORECASE,
)

# Known resume section titles — used to detect section boundaries in resumes
# that do not use Word heading styles
SECTION_TITLES_RE = re.compile(
    r"^(summary|profile|objective|highlights?"
    r"|experience|work experience|work history|professional experience|employment history"
    r"|education|academic background|academic history"
    r"|skills|technical skills|core competencies|key skills|areas of expertise"
    r"|projects?|personal projects?|side projects?"
    r"|certifications?|licenses?|credentials?"
    r"|awards?|honors?|achievements?"
    r"|publications?|research"
    r"|volunteer(ing)?|community service|community involvement"
    r"|languages?"
    r"|interests?|hobbies|activities"
    r"|references?)$",
    re.IGNORECASE,
)


def looks_like_section_title(text: str) -> bool:
    """Return True if the text looks like a resume section heading."""
    stripped = text.strip()
    if not stripped or len(stripped) > 50:
        return False
    # All-uppercase short line (common in template resumes, e.g. "EXPERIENCE")
    if stripped.isupper() and len(stripped) >= 3:
        return True
    # Exact match against known section names
    return bool(SECTION_TITLES_RE.match(stripped))


def classify_paragraphs(doc: Document) -> List[str]:
    """
    Classify every paragraph in the document into one of:
      "heading"  – section title or name heading (preserve, never edit)
      "header"   – name / contact info block at top (preserve)
      "content"  – editable body text (bullet points, descriptions, summary)
      "empty"    – blank line (preserve for spacing)

    Three-phase logic
    -----------------
    Phase 1  (before any heading)     – everything is "header" (name/contact zone)
    Phase 2  (after first heading)    – still "header" until we see a *section* heading
    Phase 3  (after a section heading) – "content", unless it looks like contact info
                                         or another section title
    """
    results = []
    first_heading_seen = False  # have we seen any heading at all?
    content_zone = False        # have we seen a section heading (h2+, or text-based)?

    for para in doc.paragraphs:
        text = para.text.strip()
        style_name = (para.style.name or "").lower()
        is_heading_style = "heading" in style_name

        if not text:
            results.append("empty")

        elif is_heading_style:
            if not first_heading_seen:
                # First heading = the candidate's name; stay in contact zone
                first_heading_seen = True
            else:
                # Second+ heading = section title; enter content zone
                content_zone = True
            results.append("heading")

        elif not first_heading_seen:
            # Phase 1: no heading seen yet – check for text-based section title
            # (resumes that never use Word heading styles, e.g., all-caps sections)
            if looks_like_section_title(text):
                first_heading_seen = True
                content_zone = True
                results.append("heading")
            else:
                results.append("header")

        elif not content_zone:
            # Phase 2: after name heading but before any section heading
            # This is the contact-info zone (address, phone, LinkedIn, etc.)
            if looks_like_section_title(text):
                content_zone = True
                results.append("heading")
            else:
                results.append("header")

        else:
            # Phase 3: content zone
            if CONTACT_RE.search(text):
                results.append("header")
            elif looks_like_section_title(text):
                # Section titles within the body (e.g., "EXPERIENCE" without heading style)
                results.append("heading")
            else:
                results.append("content")

    return results


class Section(BaseModel):
    index: int
    type: str
    text: str


class ExtractResponse(BaseModel):
    sections: List[Section]


@app.post("/api/extract-sections", response_model=ExtractResponse)
async def extract_sections(file: UploadFile = File(...)):
    data = await file.read()
    doc = Document(io.BytesIO(data))

    types = classify_paragraphs(doc)
    sections = [
        Section(index=i, type=types[i], text=doc.paragraphs[i].text)
        for i in range(len(doc.paragraphs))
    ]
    return ExtractResponse(sections=sections)


class Replacement(BaseModel):
    index: int
    new_text: str


@app.post("/api/rebuild-docx")
async def rebuild_docx(
    file: UploadFile = File(...),
    replacements: str = Form(...),
):
    data = await file.read()
    doc = Document(io.BytesIO(data))

    repl_list: List[Replacement] = [Replacement(**r) for r in json.loads(replacements)]
    repl_map = {r.index: r.new_text for r in repl_list}

    for idx, new_text in repl_map.items():
        if idx < 0 or idx >= len(doc.paragraphs):
            continue
        para = doc.paragraphs[idx]

        # Capture first run's font properties before clearing
        font_name = None
        font_size = None
        font_bold = None
        font_italic = None
        font_color = None

        if para.runs:
            first_run = para.runs[0]
            font_name = first_run.font.name
            font_size = first_run.font.size
            font_bold = first_run.font.bold
            font_italic = first_run.font.italic
            try:
                font_color = first_run.font.color.rgb if first_run.font.color.type else None
            except Exception:
                font_color = None

        # Clear all run text (keeps paragraph-level formatting: indent, bullet, spacing)
        for run in para.runs:
            run.text = ""

        # Write the new text into the first run (or a new one if none exist)
        if para.runs:
            para.runs[0].text = new_text
            run = para.runs[0]
        else:
            run = para.add_run(new_text)

        # Re-apply captured font properties
        if font_name:
            run.font.name = font_name
        if font_size:
            run.font.size = font_size
        if font_bold is not None:
            run.font.bold = font_bold
        if font_italic is not None:
            run.font.italic = font_italic
        if font_color is not None:
            try:
                run.font.color.rgb = font_color
            except Exception:
                pass

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=tailored_resume.docx"},
    )


# ── AI Complete (server-side) ──────────────────────────────────────────────────
# Calls AI providers server-side so the browser never needs CORS access.
# Used by Resume Tailoring feature.

class AICompleteRequest(BaseModel):
    provider: str
    apiKey: str
    model: str
    prompt: str
    maxTokens: int = 4096


@app.post("/api/ai/complete")
async def ai_complete(req: AICompleteRequest):
    async with httpx.AsyncClient(timeout=120.0) as client:
        if req.provider == "anthropic":
            resp = await client.post(
                "https://api.anthropic.com/v1/messages",
                headers={
                    "x-api-key": req.apiKey,
                    "anthropic-version": "2023-06-01",
                    "content-type": "application/json",
                },
                json={
                    "model": req.model,
                    "max_tokens": req.maxTokens,
                    "messages": [{"role": "user", "content": req.prompt}],
                },
            )
            if not resp.is_success:
                raise HTTPException(status_code=resp.status_code, detail=resp.text)
            text = resp.json()["content"][0]["text"]

        elif req.provider == "openai":
            resp = await client.post(
                "https://api.openai.com/v1/chat/completions",
                headers={
                    "Authorization": f"Bearer {req.apiKey}",
                    "content-type": "application/json",
                },
                json={
                    "model": req.model,
                    "messages": [{"role": "user", "content": req.prompt}],
                    "temperature": 0.3,
                },
            )
            if not resp.is_success:
                raise HTTPException(status_code=resp.status_code, detail=resp.text)
            text = resp.json()["choices"][0]["message"]["content"]

        elif req.provider == "gemini":
            url = (
                f"https://generativelanguage.googleapis.com/v1beta/models/"
                f"{req.model}:generateContent?key={req.apiKey}"
            )
            resp = await client.post(
                url,
                headers={"content-type": "application/json"},
                json={"contents": [{"parts": [{"text": req.prompt}]}]},
            )
            if not resp.is_success:
                raise HTTPException(status_code=resp.status_code, detail=resp.text)
            text = resp.json()["candidates"][0]["content"]["parts"][0]["text"]

        else:
            raise HTTPException(status_code=400, detail=f"Unknown provider: {req.provider}")

    return {"text": text}


# ── AI API Proxy ──────────────────────────────────────────────────────────────
# Forwards browser requests to Anthropic / OpenAI, bypassing browser CORS restrictions.
# Usage: set Proxy URL in Settings to http://localhost:8000/api/proxy/anthropic
#                                   or http://localhost:8000/api/proxy/openai

_PROXY_TARGETS = {
    "anthropic": "https://api.anthropic.com",
    "openai":    "https://api.openai.com",
}

_FORWARD_HEADERS = {
    "authorization", "x-api-key", "anthropic-version",
    "content-type", "anthropic-dangerous-direct-browser-access",
}


@app.api_route("/api/proxy/{provider}/{path:path}", methods=["GET", "POST", "OPTIONS"])
async def ai_proxy(provider: str, path: str, request: Request):
    if provider not in _PROXY_TARGETS:
        raise HTTPException(status_code=404, detail=f"Unknown provider: {provider}")

    target_url = f"{_PROXY_TARGETS[provider]}/{path}"
    body = await request.body()
    headers = {k: v for k, v in request.headers.items() if k.lower() in _FORWARD_HEADERS}

    async with httpx.AsyncClient(timeout=120.0) as client:
        resp = await client.request(
            method=request.method,
            url=target_url,
            headers=headers,
            content=body,
        )

    return Response(
        content=resp.content,
        status_code=resp.status_code,
        media_type=resp.headers.get("content-type", "application/json"),
    )
